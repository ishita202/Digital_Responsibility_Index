import os
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image, ImageDraw, ImageFont
import io

def set_font(run, font_name='Times New Roman', font_size=12, bold=False, italic=False):
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.bold = bold
    run.italic = italic
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), font_name)

def set_style(doc):
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    h1 = doc.styles['Heading 1']
    h1.font.name = 'Times New Roman'
    h1.font.size = Pt(16)
    h1.font.color.rgb = RGBColor(0, 0, 0)
    h1.font.bold = True
    
    h2 = doc.styles['Heading 2']
    h2.font.name = 'Times New Roman'
    h2.font.size = Pt(14)
    h2.font.color.rgb = RGBColor(0, 0, 0)
    h2.font.bold = True
    
    h3 = doc.styles['Heading 3']
    h3.font.name = 'Times New Roman'
    h3.font.size = Pt(13)
    h3.font.color.rgb = RGBColor(0, 0, 0)
    h3.font.bold = True

def create_system_architecture_diagram():
    """Create a system architecture diagram as an image"""
    width, height = 1000, 700
    img = Image.new('RGB', (width, height), 'white')
    draw = ImageDraw.Draw(img)
    
    # Try to use a font, fallback to default if not available
    try:
        font_large = ImageFont.truetype("arial.ttf", 24)
        font_medium = ImageFont.truetype("arial.ttf", 18)
        font_small = ImageFont.truetype("arial.ttf", 14)
    except:
        font_large = ImageFont.load_default()
        font_medium = ImageFont.load_default()
        font_small = ImageFont.load_default()
    
    # Colors
    blue = (91, 141, 239)  # #5B8DEF
    light_blue = (123, 163, 245)  # #7BA3F5
    coral = (255, 107, 157)  # #FF6B9D
    dark_gray = (26, 32, 44)  # #1A202C
    
    # Presentation Layer (Top)
    pres_x, pres_y = 100, 50
    pres_w, pres_h = 800, 120
    draw.rectangle([pres_x, pres_y, pres_x + pres_w, pres_y + pres_h], 
                   fill=light_blue, outline=blue, width=3)
    draw.text((pres_x + 20, pres_y + 20), "PRESENTATION LAYER", 
              fill=dark_gray, font=font_large)
    draw.text((pres_x + 20, pres_y + 60), "HTML, CSS, JavaScript, Bootstrap", 
              fill=dark_gray, font=font_medium)
    draw.text((pres_x + 20, pres_y + 90), "User Interface Components", 
              fill=dark_gray, font=font_small)
    
    # Arrow down
    arrow_y = pres_y + pres_h
    draw.line([width//2, arrow_y, width//2, arrow_y + 50], fill=blue, width=3)
    draw.polygon([(width//2, arrow_y + 50), (width//2 - 10, arrow_y + 40), 
                  (width//2 + 10, arrow_y + 40)], fill=blue)
    
    # Application Layer (Middle)
    app_x, app_y = 100, arrow_y + 50
    app_w, app_h = 800, 150
    draw.rectangle([app_x, app_y, app_x + app_w, app_y + app_h], 
                   fill=(240, 240, 255), outline=blue, width=3)
    draw.text((app_x + 20, app_y + 20), "APPLICATION LAYER", 
              fill=dark_gray, font=font_large)
    draw.text((app_x + 20, app_y + 60), "Flask Routes, Business Logic", 
              fill=dark_gray, font=font_medium)
    draw.text((app_x + 20, app_y + 90), "User Management | Quiz Engine | Analytics", 
              fill=dark_gray, font=font_small)
    draw.text((app_x + 20, app_y + 120), "Content Management | ML Integration", 
              fill=dark_gray, font=font_small)
    
    # Arrows down to bottom layers
    arrow_y2 = app_y + app_h
    draw.line([width//2 - 150, arrow_y2, width//2 - 150, arrow_y2 + 50], fill=blue, width=3)
    draw.polygon([(width//2 - 150, arrow_y2 + 50), (width//2 - 160, arrow_y2 + 40), 
                  (width//2 - 140, arrow_y2 + 40)], fill=blue)
    
    draw.line([width//2 + 150, arrow_y2, width//2 + 150, arrow_y2 + 50], fill=blue, width=3)
    draw.polygon([(width//2 + 150, arrow_y2 + 50), (width//2 + 140, arrow_y2 + 40), 
                  (width//2 + 160, arrow_y2 + 40)], fill=blue)
    
    # ML Layer (Bottom Left)
    ml_x, ml_y = 100, arrow_y2 + 50
    ml_w, ml_h = 350, 120
    draw.rectangle([ml_x, ml_y, ml_x + ml_w, ml_y + ml_h], 
                   fill=(255, 240, 245), outline=coral, width=3)
    draw.text((ml_x + 20, ml_y + 20), "ML LAYER", 
              fill=dark_gray, font=font_large)
    draw.text((ml_x + 20, ml_y + 60), "Random Forest Classifier", 
              fill=dark_gray, font=font_medium)
    draw.text((ml_x + 20, ml_y + 90), "Knowledge Level Prediction", 
              fill=dark_gray, font=font_small)
    
    # Data Layer (Bottom Right)
    data_x, data_y = 550, arrow_y2 + 50
    data_w, data_h = 350, 120
    draw.rectangle([data_x, data_y, data_x + data_w, data_y + data_h], 
                   fill=(240, 255, 240), outline=(72, 187, 120), width=3)
    draw.text((data_x + 20, data_y + 20), "DATA LAYER", 
              fill=dark_gray, font=font_large)
    draw.text((data_x + 20, data_y + 60), "SQLite Database", 
              fill=dark_gray, font=font_medium)
    draw.text((data_x + 20, data_y + 90), "SQLAlchemy ORM", 
              fill=dark_gray, font=font_small)
    
    # Save to bytes
    img_bytes = io.BytesIO()
    img.save(img_bytes, format='PNG')
    img_bytes.seek(0)
    return img_bytes

def create_er_diagram():
    """Create a database ER diagram as an image"""
    width, height = 1200, 800
    img = Image.new('RGB', (width, height), 'white')
    draw = ImageDraw.Draw(img)
    
    try:
        font_large = ImageFont.truetype("arial.ttf", 20)
        font_medium = ImageFont.truetype("arial.ttf", 16)
        font_small = ImageFont.truetype("arial.ttf", 12)
    except:
        font_large = ImageFont.load_default()
        font_medium = ImageFont.load_default()
        font_small = ImageFont.load_default()
    
    blue = (91, 141, 239)
    dark_gray = (26, 32, 44)
    green = (72, 187, 120)
    
    # User Entity (Top Center)
    user_x, user_y = 500, 50
    user_w, user_h = 200, 180
    draw.rectangle([user_x, user_y, user_x + user_w, user_y + user_h], 
                   fill=(240, 240, 255), outline=blue, width=3)
    draw.text((user_x + 20, user_y + 10), "USER", fill=dark_gray, font=font_large)
    draw.text((user_x + 20, user_y + 40), "id (PK)", fill=dark_gray, font=font_small)
    draw.text((user_x + 20, user_y + 60), "username", fill=dark_gray, font=font_small)
    draw.text((user_x + 20, user_y + 80), "email", fill=dark_gray, font=font_small)
    draw.text((user_x + 20, user_y + 100), "password_hash", fill=dark_gray, font=font_small)
    draw.text((user_x + 20, user_y + 120), "is_admin", fill=dark_gray, font=font_small)
    draw.text((user_x + 20, user_y + 140), "age_range, gender", fill=dark_gray, font=font_small)
    draw.text((user_x + 20, user_y + 160), "academic_stream, etc.", fill=dark_gray, font=font_small)
    
    # QuizAttempt (Bottom Left)
    qa_x, qa_y = 100, 350
    qa_w, qa_h = 200, 180
    draw.rectangle([qa_x, qa_y, qa_x + qa_w, qa_y + qa_h], 
                   fill=(255, 240, 245), outline=(255, 107, 157), width=3)
    draw.text((qa_x + 20, qa_y + 10), "QUIZ ATTEMPT", fill=dark_gray, font=font_large)
    draw.text((qa_x + 20, qa_y + 40), "id (PK)", fill=dark_gray, font=font_small)
    draw.text((qa_x + 20, qa_y + 60), "user_id (FK)", fill=dark_gray, font=font_small)
    draw.text((qa_x + 20, qa_y + 80), "quiz_type", fill=dark_gray, font=font_small)
    draw.text((qa_x + 20, qa_y + 100), "score, percentage", fill=dark_gray, font=font_small)
    draw.text((qa_x + 20, qa_y + 120), "time_taken", fill=dark_gray, font=font_small)
    draw.text((qa_x + 20, qa_y + 140), "answers (JSON)", fill=dark_gray, font=font_small)
    draw.text((qa_x + 20, qa_y + 160), "completed_at", fill=dark_gray, font=font_small)
    
    # UserActivity (Bottom Right)
    ua_x, ua_y = 900, 350
    ua_w, ua_h = 200, 150
    draw.rectangle([ua_x, ua_y, ua_x + ua_w, ua_y + ua_h], 
                   fill=(255, 250, 240), outline=(255, 184, 77), width=3)
    draw.text((ua_x + 20, ua_y + 10), "USER ACTIVITY", fill=dark_gray, font=font_large)
    draw.text((ua_x + 20, ua_y + 40), "id (PK)", fill=dark_gray, font=font_small)
    draw.text((ua_x + 20, ua_y + 60), "user_id (FK)", fill=dark_gray, font=font_small)
    draw.text((ua_x + 20, ua_y + 80), "activity_type", fill=dark_gray, font=font_small)
    draw.text((ua_x + 20, ua_y + 100), "description", fill=dark_gray, font=font_small)
    draw.text((ua_x + 20, ua_y + 120), "created_at", fill=dark_gray, font=font_small)
    
    # QuizQuestion (Top Left)
    qq_x, qq_y = 100, 50
    qq_w, qq_h = 200, 200
    draw.rectangle([qq_x, qq_y, qq_x + qq_w, qq_y + qq_h], 
                   fill=(240, 255, 240), outline=green, width=3)
    draw.text((qq_x + 20, qq_y + 10), "QUIZ QUESTION", fill=dark_gray, font=font_large)
    draw.text((qq_x + 20, qq_y + 40), "id (PK)", fill=dark_gray, font=font_small)
    draw.text((qq_x + 20, qq_y + 60), "question_text", fill=dark_gray, font=font_small)
    draw.text((qq_x + 20, qq_y + 80), "option_a, b, c, d", fill=dark_gray, font=font_small)
    draw.text((qq_x + 20, qq_y + 100), "correct_answer", fill=dark_gray, font=font_small)
    draw.text((qq_x + 20, qq_y + 120), "category, quiz_type", fill=dark_gray, font=font_small)
    draw.text((qq_x + 20, qq_y + 140), "explanation", fill=dark_gray, font=font_small)
    draw.text((qq_x + 20, qq_y + 160), "difficulty, time_limit", fill=dark_gray, font=font_small)
    draw.text((qq_x + 20, qq_y + 180), "created_at", fill=dark_gray, font=font_small)
    
    # LearningResource (Top Right)
    lr_x, lr_y = 900, 50
    lr_w, lr_h = 200, 150
    draw.rectangle([lr_x, lr_y, lr_x + lr_w, lr_y + lr_h], 
                   fill=(255, 255, 240), outline=(255, 212, 77), width=3)
    draw.text((lr_x + 20, lr_y + 10), "LEARNING RESOURCE", fill=dark_gray, font=font_large)
    draw.text((lr_x + 20, lr_y + 40), "id (PK)", fill=dark_gray, font=font_small)
    draw.text((lr_x + 20, lr_y + 60), "title, description", fill=dark_gray, font=font_small)
    draw.text((lr_x + 20, lr_y + 80), "url", fill=dark_gray, font=font_small)
    draw.text((lr_x + 20, lr_y + 100), "category, resource_type", fill=dark_gray, font=font_small)
    draw.text((lr_x + 20, lr_y + 120), "created_at", fill=dark_gray, font=font_small)
    
    # QuizType (Bottom Center)
    qt_x, qt_y = 500, 550
    qt_w, qt_h = 200, 120
    draw.rectangle([qt_x, qt_y, qt_x + qt_w, qt_y + qt_h], 
                   fill=(245, 240, 255), outline=(147, 51, 234), width=3)
    draw.text((qt_x + 20, qt_y + 10), "QUIZ TYPE", fill=dark_gray, font=font_large)
    draw.text((qt_x + 20, qt_y + 40), "id (PK)", fill=dark_gray, font=font_small)
    draw.text((qt_x + 20, qt_y + 60), "name, description", fill=dark_gray, font=font_small)
    draw.text((qt_x + 20, qt_y + 80), "time_limit, difficulty", fill=dark_gray, font=font_small)
    draw.text((qt_x + 20, qt_y + 100), "question_count", fill=dark_gray, font=font_small)
    
    # Relationships
    # User -> QuizAttempt (1 to Many)
    draw.line([user_x + 100, user_y + user_h, qa_x + 100, qa_y], fill=blue, width=2)
    draw.text((user_x + 50, user_y + user_h + 20), "1", fill=blue, font=font_medium)
    draw.text((qa_x + 50, qa_y - 20), "Many", fill=blue, font=font_medium)
    
    # User -> UserActivity (1 to Many)
    draw.line([user_x + user_w, user_y + 100, ua_x, ua_y + 50], fill=blue, width=2)
    draw.text((user_x + user_w + 10, user_y + 80), "1", fill=blue, font=font_medium)
    draw.text((ua_x - 50, ua_y + 30), "Many", fill=blue, font=font_medium)
    
    # QuizQuestion -> QuizAttempt (Many to Many through answers JSON)
    draw.line([qq_x + qq_w, qq_y + 100, qa_x, qa_y + 50], fill=green, width=2)
    draw.text((qq_x + qq_w + 10, qq_y + 80), "Many", fill=green, font=font_medium)
    draw.text((qa_x - 50, qa_y + 30), "Many", fill=green, font=font_medium)
    
    img_bytes = io.BytesIO()
    img.save(img_bytes, format='PNG')
    img_bytes.seek(0)
    return img_bytes

def add_cover_page(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    run = p.add_run("A PROJECT REPORT\nON\n")
    set_font(run, font_size=16, bold=True)
    
    run = p.add_run("\nDIGITAL AWARENESS PLATFORM\n")
    set_font(run, font_size=20, bold=True)
    
    run = p.add_run("\nSUBMITTED IN PARTIAL FULFILLMENT OF THE REQUIREMENTS FOR THE AWARD OF THE DEGREE OF\n\n")
    set_font(run, font_size=14)
    
    run = p.add_run("BACHELOR OF TECHNOLOGY\nIN\nCOMPUTER SCIENCE AND ENGINEERING\n\n")
    set_font(run, font_size=16, bold=True)
    
    run = p.add_run("By\n\n")
    set_font(run, font_size=14)
    
    run = p.add_run("[STUDENT NAME]\n(Roll No: XXXXXXX)\n\n")
    set_font(run, font_size=14, bold=True)
    
    run = p.add_run("Under the Guidance of\n\n")
    set_font(run, font_size=14)
    
    run = p.add_run("[GUIDE NAME]\n[Designation]\n\n")
    set_font(run, font_size=14, bold=True)
    
    run = p.add_run("DEPARTMENT OF COMPUTER SCIENCE & ENGINEERING\n[INSTITUTE NAME]\n[CITY, STATE - PIN CODE]\n2024-2025")
    set_font(run, font_size=16, bold=True)

    doc.add_page_break()

def add_preliminary_pages(doc):
    doc.add_heading('ACKNOWLEDGEMENT', level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    text = ("I would like to express my gratitude to my project guide [Guide Name] for their guidance and support. "
            "I am also thankful to the Head of Department and the faculty members for their encouragement. "
            "Finally, I would like to thank my friends and family for their continuous support throughout this project.")
    run = p.add_run(text)
    set_font(run)
    doc.add_page_break()

    doc.add_heading('DECLARATION', level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    text = ("I hereby declare that the project report entitled 'Digital Awareness Platform' submitted by me to [Institute Name] "
            "in partial fulfillment of the requirement for the award of the degree of Bachelor of Technology in "
            "Computer Science and Engineering is a record of my own work carried out under the supervision of [Guide Name].")
    run = p.add_run(text)
    set_font(run)
    
    p = doc.add_paragraph("\n\nDate: ____________\nPlace: ____________")
    set_font(p.add_run())
    
    p = doc.add_paragraph("\n\n(Signature of the Student)\n[Student Name]")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_font(p.add_run())
    doc.add_page_break()

def clean_md_text(text):
    text = text.replace('**', '').replace('*', '')
    text = text.replace('`', '')
    return text.strip()

def analyze_document_structure(lines):
    toc_entries = []
    tables = []
    figures = []
    
    for line in lines:
        line = line.strip()
        
        if line.startswith('# '):
            text = clean_md_text(line[2:])
            if text.upper() not in ["DIGITAL AWARENESS PLATFORM", "TABLE OF CONTENTS"]:
                toc_entries.append((1, text))
        elif line.startswith('## '):
            text = clean_md_text(line[3:])
            if text.upper() not in ["TABLE OF CONTENTS", "ABSTRACT", "ACKNOWLEDGEMENT", "DECLARATION", "REFERENCES", "APPENDICES"]:
                if re.match(r'^\d+\.\s', text) or text in ["Introduction", "Literature Review", "System Analysis", "Implementation", "Testing"]:
                    toc_entries.append((1, text))
                else:
                    toc_entries.append((2, text))
            elif text.upper() in ["ABSTRACT", "REFERENCES", "APPENDICES"]:
                toc_entries.append((1, text))
                
        elif line.startswith('### '):
            text = clean_md_text(line[4:])
            toc_entries.append((2, text))
        elif line.startswith('#### '):
            text = clean_md_text(line[5:])
            toc_entries.append((3, text))
            
        if re.match(r'\*\*Table \d+:', line) or re.match(r'Table \d+:', line):
            tables.append(clean_md_text(line))
            
    # Add detected figures
    figures.append("Figure 1: System Architecture Diagram")
    figures.append("Figure 2: Database Entity Relationship Diagram")
    
    return toc_entries, tables, figures

def add_toc(doc, entries):
    doc.add_heading('TABLE OF CONTENTS', level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    for level, text in entries:
        p = doc.add_paragraph()
        tab_stops = p.paragraph_format.tab_stops
        tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
        
        indent = Inches(0.2) * (level - 1)
        p.paragraph_format.left_indent = indent
        
        run = p.add_run(text)
        set_font(run, bold=(level==1))
    
    doc.add_page_break()

def add_list_of_figures(doc, figures):
    doc.add_heading('LIST OF FIGURES', level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    if not figures:
        doc.add_paragraph("No figures found.")
    else:
        for fig in figures:
            p = doc.add_paragraph()
            tab_stops = p.paragraph_format.tab_stops
            tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
            run = p.add_run(fig)
            set_font(run)
            p.add_run('\t')
            
    doc.add_page_break()

def add_list_of_tables(doc, tables):
    doc.add_heading('LIST OF TABLES', level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    if not tables:
         tables = [
             "Table 1: User",
             "Table 2: QuizQuestion", 
             "Table 3: QuizAttempt",
             "Table 4: UserActivity",
             "Table 5: LearningResource",
             "Table 6: QuizType",
             "Table 7: Backend Technologies",
             "Table 8: Machine Learning Stack",
             "Table 9: Frontend Technologies",
             "Table 10: Functional Testing Results",
             "Table 11: Performance Metrics"
         ]
         
    for tbl in tables:
        p = doc.add_paragraph()
        tab_stops = p.paragraph_format.tab_stops
        tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
        run = p.add_run(tbl)
        set_font(run)
        p.add_run('\t')
            
    doc.add_page_break()

def parse_md_and_add_to_doc(doc, md_file_path):
    with open(md_file_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    toc_entries, tables, figures = analyze_document_structure(lines)
    
    abstract_lines = []
    in_abstract = False
    for line in lines:
        if "## ABSTRACT" in line.upper():
            in_abstract = True
            continue
        if in_abstract and line.startswith("#"):
            in_abstract = False
        if in_abstract:
            abstract_lines.append(line)
            
    if abstract_lines:
        doc.add_heading('ABSTRACT', level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
        for line in abstract_lines:
            text = clean_md_text(line.strip())
            if text:
                if text == "---": continue
                p = doc.add_paragraph(text)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                set_font(p.runs[0])
        doc.add_page_break()
        
    add_toc(doc, toc_entries)
    add_list_of_figures(doc, figures)
    add_list_of_tables(doc, tables)
    
    code_block = False
    code_content = []
    
    start_processing = False
    skip_mode = False 
    last_was_page_break = True
    diagram_added = False
    er_added = False
    
    for line in lines:
        line_stripped = line.strip()
        
        if "## ABSTRACT" in line.upper():
            skip_mode = True
            continue
        if "## TABLE OF CONTENTS" in line.upper():
            skip_mode = True
            continue
            
        if "## 1. INTRODUCTION" in line.upper() or "## 2. INTRODUCTION" in line.upper():
            start_processing = True
            skip_mode = False
            last_was_page_break = False
            
        if not start_processing:
            continue
            
        if skip_mode:
            if line.startswith('## ') and "ABSTRACT" not in line.upper() and "TABLE OF CONTENTS" not in line.upper():
                 skip_mode = False
            else:
                 continue

        # Add System Architecture Diagram when we reach that section
        if "### 3.2.1 Overall Architecture" in line or ("System Architecture" in line and "###" in line and not diagram_added):
            if not diagram_added:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                img_bytes = create_system_architecture_diagram()
                run.add_picture(img_bytes, width=Inches(6))
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run("Fig. 1 System Architecture Diagram")
                set_font(run, bold=True)
                diagram_added = True
                last_was_page_break = False
                continue

        # Add ER Diagram when we reach database design section
        if ("### 3.3.1 Entity Relationship Model" in line or 
            ("Entity Relationship" in line and "###" in line and not er_added)):
            if not er_added:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                img_bytes = create_er_diagram()
                run.add_picture(img_bytes, width=Inches(6))
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run("Fig. 2 Database Entity Relationship Diagram")
                set_font(run, bold=True)
                er_added = True
                last_was_page_break = False
                continue

        # Code blocks
        if line.startswith('```'):
            if code_block:
                p = doc.add_paragraph()
                run = p.add_run('\n'.join(code_content))
                run.font.name = 'Courier New'
                run.font.size = Pt(10)
                code_content = []
                code_block = False
                last_was_page_break = False
            else:
                code_block = True
            continue
            
        if code_block:
            code_content.append(line.rstrip())
            continue
            
        # Headers
        if line.startswith('# '):
            text = clean_md_text(line[2:])
            if text.upper() == "DIGITAL AWARENESS PLATFORM": continue 
            
            if not last_was_page_break:
                doc.add_page_break()
            h = doc.add_heading(text, level=1)
            h.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for r in h.runs:
                r.font.color.rgb = RGBColor(0,0,0)
                r.font.name = 'Times New Roman'
                r.font.size = Pt(16)
            last_was_page_break = False
        
        elif line.startswith('## '):
            text = clean_md_text(line[3:])
            if text.upper() in ["ABSTRACT", "TABLE OF CONTENTS"]: continue
            
            is_chapter = re.match(r'^\d+\.\s+[A-Z\s]+$', text) or text.upper() in ["REFERENCES", "APPENDICES"]
            
            if is_chapter:
                if not last_was_page_break:
                    doc.add_page_break()
                h = doc.add_heading(text, level=1)
                for r in h.runs:
                    r.font.color.rgb = RGBColor(0,0,0)
                    r.font.name = 'Times New Roman'
                    r.font.size = Pt(16)
            else:
                h = doc.add_heading(text, level=2)
                for r in h.runs:
                    r.font.color.rgb = RGBColor(0,0,0)
                    r.font.name = 'Times New Roman'
                    r.font.size = Pt(14)
            last_was_page_break = False
                    
        elif line.startswith('### '):
            text = clean_md_text(line[4:])
            h = doc.add_heading(text, level=2)
            for r in h.runs:
                r.font.color.rgb = RGBColor(0,0,0)
                r.font.name = 'Times New Roman'
                r.font.size = Pt(14)
            last_was_page_break = False
                
        elif line.startswith('#### '):
            text = clean_md_text(line[5:])
            h = doc.add_heading(text, level=3)
            for r in h.runs:
                r.font.color.rgb = RGBColor(0,0,0)
                r.font.name = 'Times New Roman'
                r.font.size = Pt(12)
                r.font.bold = True
            last_was_page_break = False
                
        elif line_stripped.startswith('- '):
            text = clean_md_text(line_stripped[2:])
            p = doc.add_paragraph(text, style='List Bullet')
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            for r in p.runs:
                set_font(r)
            last_was_page_break = False
                
        elif re.match(r'^\d+\.\s', line_stripped):
            text = clean_md_text(re.sub(r'^\d+\.\s', '', line_stripped))
            p = doc.add_paragraph(text, style='List Number')
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            for r in p.runs:
                set_font(r)
            last_was_page_break = False
        
        elif line_stripped.startswith('|'):
            p = doc.add_paragraph(line_stripped)
            p.style = 'No Spacing'
            run = p.runs[0]
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
            last_was_page_break = False
            
        else:
            if not line_stripped:
                continue
            text = clean_md_text(line_stripped)
            if text == "---": 
                if not last_was_page_break:
                    doc.add_page_break()
                    last_was_page_break = True
                continue
                
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            run = p.add_run(text)
            set_font(run)
            last_was_page_break = False

def main():
    doc = Document()
    set_style(doc)
    
    add_cover_page(doc)
    add_preliminary_pages(doc)
    
    parse_md_and_add_to_doc(doc, 'PROJECT_REPORT_STANDARD.md')
    
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    output_path = 'Digital_Awareness_Platform_Project_Report_Final.docx'
    try:
        doc.save(output_path)
        print(f"Report generated: {output_path}")
    except PermissionError:
        print(f"Error: Could not save to {output_path}. Please close the file if it is open.")
        output_path = 'Digital_Awareness_Platform_Project_Report_Final_v2.docx'
        doc.save(output_path)
        print(f"Report generated: {output_path}")

if __name__ == "__main__":
    main()

