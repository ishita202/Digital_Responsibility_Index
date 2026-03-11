import os
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image, ImageDraw, ImageFont
import io

def set_font(run, font_name='Times New Roman', font_size=12, bold=False, italic=False):
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.bold = bold
    run.italic = italic
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), font_name)

def set_style(doc):
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    h1 = doc.styles['Heading 1']
    h1.font.name = 'Times New Roman'
    h1.font.size = Pt(16)
    h1.font.color.rgb = RGBColor(0, 0, 0)
    h1.font.bold = True
    
    h2 = doc.styles['Heading 2']
    h2.font.name = 'Times New Roman'
    h2.font.size = Pt(14)
    h2.font.color.rgb = RGBColor(0, 0, 0)
    h2.font.bold = True
    
    h3 = doc.styles['Heading 3']
    h3.font.name = 'Times New Roman'
    h3.font.size = Pt(13)
    h3.font.color.rgb = RGBColor(0, 0, 0)
    h3.font.bold = True

def create_system_architecture_diagram():
    """Create a comprehensive system architecture diagram showing all components"""
    width, height = 1400, 900
    img = Image.new('RGB', (width, height), 'white')
    draw = ImageDraw.Draw(img)
    
    try:
        font_large = ImageFont.truetype("arial.ttf", 22)
        font_medium = ImageFont.truetype("arial.ttf", 16)
        font_small = ImageFont.truetype("arial.ttf", 12)
    except:
        font_large = ImageFont.load_default()
        font_medium = ImageFont.load_default()
        font_small = ImageFont.load_default()
    
    blue = (91, 141, 239)
    light_blue = (123, 163, 245)
    coral = (255, 107, 157)
    dark_gray = (26, 32, 44)
    green = (72, 187, 120)
    orange = (255, 184, 77)
    purple = (147, 51, 234)
    
    # Helper function for dashed lines
    def draw_dashed_line(x1, y1, x2, y2, color, width=2, dash_length=5):
        if x1 == x2:  # Vertical line
            y = y1
            while y < y2:
                draw.line([x1, y, x1, min(y + dash_length, y2)], fill=color, width=width)
                y += dash_length * 2
        else:  # Horizontal line
            x = x1
            while x < x2:
                draw.line([x, y1, min(x + dash_length, x2), y1], fill=color, width=width)
                x += dash_length * 2
    
    # Users Layer (Top)
    users_x, users_y = 50, 30
    users_w, users_h = 1300, 80
    draw.rectangle([users_x, users_y, users_x + users_w, users_y + users_h], 
                   fill=(255, 250, 240), outline=orange, width=3)
    draw.text((users_x + 20, users_y + 15), "USERS", 
              fill=dark_gray, font=font_large)
    draw.text((users_x + 20, users_y + 45), "Students | Professionals | Administrators | Content Managers", 
              fill=dark_gray, font=font_medium)
    draw.text((users_x + 600, users_y + 15), "Access via: Web Browsers (Desktop, Tablet, Mobile)", 
              fill=dark_gray, font=font_small)
    
    # Arrow down to Frontend
    arrow_y1 = users_y + users_h
    draw.line([width//2, arrow_y1, width//2, arrow_y1 + 40], fill=blue, width=3)
    draw.polygon([(width//2, arrow_y1 + 40), (width//2 - 10, arrow_y1 + 30), 
                  (width//2 + 10, arrow_y1 + 30)], fill=blue)
    
    # Frontend (Web Application) Layer
    frontend_x, frontend_y = 50, arrow_y1 + 40
    frontend_w, frontend_h = 600, 140
    draw.rectangle([frontend_x, frontend_y, frontend_x + frontend_w, frontend_y + frontend_h], 
                   fill=light_blue, outline=blue, width=3)
    draw.text((frontend_x + 20, frontend_y + 15), "FRONTEND (WEB APPLICATION)", 
              fill=dark_gray, font=font_large)
    draw.text((frontend_x + 20, frontend_y + 50), "HTML5 | CSS3 | JavaScript | Bootstrap 5", 
              fill=dark_gray, font=font_medium)
    draw.text((frontend_x + 20, frontend_y + 75), "User Interface Templates | Responsive Design", 
              fill=dark_gray, font=font_small)
    draw.text((frontend_x + 20, frontend_y + 95), "Quiz Interface | Dashboard | Admin Panel", 
              fill=dark_gray, font=font_small)
    draw.text((frontend_x + 20, frontend_y + 115), "AJAX | RESTful API Calls", 
              fill=dark_gray, font=font_small)
    
    # Backend Server Layer (Right of Frontend)
    backend_x, backend_y = 700, arrow_y1 + 40
    backend_w, backend_h = 650, 140
    draw.rectangle([backend_x, backend_y, backend_x + backend_w, backend_y + backend_h], 
                   fill=(240, 240, 255), outline=blue, width=3)
    draw.text((backend_x + 20, backend_y + 15), "BACKEND SERVER (FLASK APPLICATION)", 
              fill=dark_gray, font=font_large)
    draw.text((backend_x + 20, backend_y + 50), "Python Flask | Route Handlers | Business Logic", 
              fill=dark_gray, font=font_medium)
    draw.text((backend_x + 20, backend_y + 75), "Authentication | Quiz Engine | ML Integration", 
              fill=dark_gray, font=font_small)
    draw.text((backend_x + 20, backend_y + 95), "Admin Dashboard | Content Management | API Endpoints", 
              fill=dark_gray, font=font_small)
    draw.text((backend_x + 20, backend_y + 115), "Session Management | Request Processing", 
              fill=dark_gray, font=font_small)
    
    # Arrow from Frontend to Backend
    arrow_mid = frontend_x + frontend_w
    arrow_end = backend_x
    arrow_y_mid = frontend_y + frontend_h//2
    draw.line([arrow_mid, arrow_y_mid, arrow_end, arrow_y_mid], fill=blue, width=3)
    draw.polygon([(arrow_end, arrow_y_mid), (arrow_end - 10, arrow_y_mid - 5), 
                  (arrow_end - 10, arrow_y_mid + 5)], fill=blue)
    draw.text((arrow_mid + 20, arrow_y_mid - 10), "HTTP/HTTPS", fill=blue, font=font_small)
    
    # Arrows down from Backend
    arrow_y2 = backend_y + backend_h
    draw.line([backend_x + backend_w//2 - 100, arrow_y2, backend_x + backend_w//2 - 100, arrow_y2 + 40], fill=blue, width=3)
    draw.polygon([(backend_x + backend_w//2 - 100, arrow_y2 + 40), 
                  (backend_x + backend_w//2 - 110, arrow_y2 + 30), 
                  (backend_x + backend_w//2 - 90, arrow_y2 + 30)], fill=blue)
    
    draw.line([backend_x + backend_w//2, arrow_y2, backend_x + backend_w//2, arrow_y2 + 40], fill=blue, width=3)
    draw.polygon([(backend_x + backend_w//2, arrow_y2 + 40), 
                  (backend_x + backend_w//2 - 10, arrow_y2 + 30), 
                  (backend_x + backend_w//2 + 10, arrow_y2 + 30)], fill=blue)
    
    draw.line([backend_x + backend_w//2 + 100, arrow_y2, backend_x + backend_w//2 + 100, arrow_y2 + 40], fill=blue, width=3)
    draw.polygon([(backend_x + backend_w//2 + 100, arrow_y2 + 40), 
                  (backend_x + backend_w//2 + 90, arrow_y2 + 30), 
                  (backend_x + backend_w//2 + 110, arrow_y2 + 30)], fill=blue)
    
    # Database Layer (Bottom Left)
    db_x, db_y = 50, arrow_y2 + 40
    db_w, db_h = 400, 140
    draw.rectangle([db_x, db_y, db_x + db_w, db_y + db_h], 
                   fill=(240, 255, 240), outline=green, width=3)
    draw.text((db_x + 20, db_y + 15), "DATABASE (SQLITE)", 
              fill=dark_gray, font=font_large)
    draw.text((db_x + 20, db_y + 45), "SQLAlchemy ORM | Data Persistence", 
              fill=dark_gray, font=font_medium)
    draw.text((db_x + 20, db_y + 70), "User | QuizQuestion | QuizAttempt", 
              fill=dark_gray, font=font_small)
    draw.text((db_x + 20, db_y + 90), "UserActivity | LearningResource | QuizType", 
              fill=dark_gray, font=font_small)
    draw.text((db_x + 20, db_y + 110), "Foreign Keys | Data Integrity | Query Optimization", 
              fill=dark_gray, font=font_small)
    
    # ML Model Layer (Bottom Center)
    ml_x, ml_y = 500, arrow_y2 + 40
    ml_w, ml_h = 400, 140
    draw.rectangle([ml_x, ml_y, ml_x + ml_w, ml_y + ml_h], 
                   fill=(255, 240, 245), outline=coral, width=3)
    draw.text((ml_x + 20, ml_y + 15), "MACHINE LEARNING MODEL", 
              fill=dark_gray, font=font_large)
    draw.text((ml_x + 20, ml_y + 45), "Random Forest Classifier | scikit-learn", 
              fill=dark_gray, font=font_medium)
    draw.text((ml_x + 20, ml_y + 70), "Knowledge Level Prediction (Low/Medium/High)", 
              fill=dark_gray, font=font_small)
    draw.text((ml_x + 20, ml_y + 90), "Personalized Recommendations | Real-time Inference", 
              fill=dark_gray, font=font_small)
    draw.text((ml_x + 20, ml_y + 110), "Model Persistence (ml_model.pkl)", 
              fill=dark_gray, font=font_small)
    
    # APIs / External Services Layer (Bottom Right)
    api_x, api_y = 950, arrow_y2 + 40
    api_w, api_h = 400, 140
    draw.rectangle([api_x, api_y, api_x + api_w, api_y + api_h], 
                   fill=(255, 255, 240), outline=orange, width=3)
    draw.text((api_x + 20, api_y + 15), "APIs / EXTERNAL SERVICES", 
              fill=dark_gray, font=font_large)
    draw.text((api_x + 20, api_y + 45), "Google Sheets API | Survey Data Import", 
              fill=dark_gray, font=font_medium)
    draw.text((api_x + 20, api_y + 70), "Data Synchronization | CSV Export", 
              fill=dark_gray, font=font_small)
    draw.text((api_x + 20, api_y + 90), "Service Account Authentication", 
              fill=dark_gray, font=font_small)
    draw.text((api_x + 20, api_y + 110), "Future: Email Services | Analytics APIs", 
              fill=dark_gray, font=font_small)
    
    # Cloud / Hosting Layer (Bottom)
    cloud_y = db_y + db_h + 30
    cloud_x, cloud_w, cloud_h = 50, 1300, 100
    draw.rectangle([cloud_x, cloud_y, cloud_x + cloud_w, cloud_y + cloud_h], 
                   fill=(245, 240, 255), outline=purple, width=3)
    draw.text((cloud_x + 20, cloud_y + 15), "CLOUD / HOSTING INFRASTRUCTURE", 
              fill=dark_gray, font=font_large)
    draw.text((cloud_x + 20, cloud_y + 45), "Heroku | AWS | Google Cloud Platform | Local Development", 
              fill=dark_gray, font=font_medium)
    draw.text((cloud_x + 20, cloud_y + 70), "Web Server | Database Storage | Static File Serving | Scalability", 
              fill=dark_gray, font=font_small)
    
    # Arrows from components to Cloud (dashed lines)
    draw_dashed_line(db_x + db_w//2, db_y + db_h, db_x + db_w//2, cloud_y, purple, 2)
    draw_dashed_line(ml_x + ml_w//2, ml_y + ml_h, ml_x + ml_w//2, cloud_y, purple, 2)
    draw_dashed_line(api_x + api_w//2, api_y + api_h, api_x + api_w//2, cloud_y, purple, 2)
    
    img_bytes = io.BytesIO()
    img.save(img_bytes, format='PNG')
    img_bytes.seek(0)
    return img_bytes

def create_ml_pipeline_diagram():
    """Create ML pipeline diagram"""
    width, height = 1200, 600
    img = Image.new('RGB', (width, height), 'white')
    draw = ImageDraw.Draw(img)
    
    try:
        font_large = ImageFont.truetype("arial.ttf", 20)
        font_medium = ImageFont.truetype("arial.ttf", 16)
        font_small = ImageFont.truetype("arial.ttf", 12)
    except:
        font_large = ImageFont.load_default()
        font_medium = ImageFont.load_default()
        font_small = ImageFont.load_default()
    
    blue = (91, 141, 239)
    dark_gray = (26, 32, 44)
    green = (72, 187, 120)
    orange = (255, 184, 77)
    coral = (255, 107, 157)
    
    stages = [
        ("Data Collection", 100, 100, 180, 100),
        ("Preprocessing", 320, 100, 180, 100),
        ("Model Training", 540, 100, 180, 100),
        ("Model Persistence", 760, 100, 180, 100),
        ("Prediction", 980, 100, 180, 100),
    ]
    
    for i, (label, x, y, w, h) in enumerate(stages):
        color = [blue, green, orange, (147, 51, 234), coral][i]
        draw.rectangle([x, y, x + w, y + h], fill=(*[c//2 + 128 for c in color],), outline=color, width=2)
        draw.text((x + 10, y + 30), label, fill=dark_gray, font=font_medium)
        
        if i < len(stages) - 1:
            arrow_x = x + w
            draw.line([arrow_x, y + h//2, arrow_x + 40, y + h//2], fill=blue, width=2)
            draw.polygon([(arrow_x + 40, y + h//2), (arrow_x + 30, y + h//2 - 5), 
                          (arrow_x + 30, y + h//2 + 5)], fill=blue)
    
    img_bytes = io.BytesIO()
    img.save(img_bytes, format='PNG')
    img_bytes.seek(0)
    return img_bytes

def create_er_diagram():
    """Create a database ER diagram as an image"""
    width, height = 1200, 800
    img = Image.new('RGB', (width, height), 'white')
    draw = ImageDraw.Draw(img)
    
    try:
        font_large = ImageFont.truetype("arial.ttf", 20)
        font_medium = ImageFont.truetype("arial.ttf", 16)
        font_small = ImageFont.truetype("arial.ttf", 12)
    except:
        font_large = ImageFont.load_default()
        font_medium = ImageFont.load_default()
        font_small = ImageFont.load_default()
    
    blue = (91, 141, 239)
    dark_gray = (26, 32, 44)
    green = (72, 187, 120)
    
    # User Entity (Top Center)
    user_x, user_y = 500, 50
    user_w, user_h = 200, 180
    draw.rectangle([user_x, user_y, user_x + user_w, user_y + user_h], 
                   fill=(240, 240, 255), outline=blue, width=3)
    draw.text((user_x + 20, user_y + 10), "USER", fill=dark_gray, font=font_large)
    draw.text((user_x + 20, user_y + 40), "id (PK)", fill=dark_gray, font=font_small)
    draw.text((user_x + 20, user_y + 60), "username", fill=dark_gray, font=font_small)
    draw.text((user_x + 20, user_y + 80), "email", fill=dark_gray, font=font_small)
    draw.text((user_x + 20, user_y + 100), "password_hash", fill=dark_gray, font=font_small)
    draw.text((user_x + 20, user_y + 120), "is_admin", fill=dark_gray, font=font_small)
    draw.text((user_x + 20, user_y + 140), "age_range, gender", fill=dark_gray, font=font_small)
    draw.text((user_x + 20, user_y + 160), "academic_stream, etc.", fill=dark_gray, font=font_small)
    
    # QuizAttempt (Bottom Left)
    qa_x, qa_y = 100, 350
    qa_w, qa_h = 200, 180
    draw.rectangle([qa_x, qa_y, qa_x + qa_w, qa_y + qa_h], 
                   fill=(255, 240, 245), outline=(255, 107, 157), width=3)
    draw.text((qa_x + 20, qa_y + 10), "QUIZ ATTEMPT", fill=dark_gray, font=font_large)
    draw.text((qa_x + 20, qa_y + 40), "id (PK)", fill=dark_gray, font=font_small)
    draw.text((qa_x + 20, qa_y + 60), "user_id (FK)", fill=dark_gray, font=font_small)
    draw.text((qa_x + 20, qa_y + 80), "quiz_type", fill=dark_gray, font=font_small)
    draw.text((qa_x + 20, qa_y + 100), "score, percentage", fill=dark_gray, font=font_small)
    draw.text((qa_x + 20, qa_y + 120), "time_taken", fill=dark_gray, font=font_small)
    draw.text((qa_x + 20, qa_y + 140), "answers (JSON)", fill=dark_gray, font=font_small)
    draw.text((qa_x + 20, qa_y + 160), "completed_at", fill=dark_gray, font=font_small)
    
    # UserActivity (Bottom Right)
    ua_x, ua_y = 900, 350
    ua_w, ua_h = 200, 150
    draw.rectangle([ua_x, ua_y, ua_x + ua_w, ua_y + ua_h], 
                   fill=(255, 250, 240), outline=(255, 184, 77), width=3)
    draw.text((ua_x + 20, ua_y + 10), "USER ACTIVITY", fill=dark_gray, font=font_large)
    draw.text((ua_x + 20, ua_y + 40), "id (PK)", fill=dark_gray, font=font_small)
    draw.text((ua_x + 20, ua_y + 60), "user_id (FK)", fill=dark_gray, font=font_small)
    draw.text((ua_x + 20, ua_y + 80), "activity_type", fill=dark_gray, font=font_small)
    draw.text((ua_x + 20, ua_y + 100), "description", fill=dark_gray, font=font_small)
    draw.text((ua_x + 20, ua_y + 120), "created_at", fill=dark_gray, font=font_small)
    
    # QuizQuestion (Top Left)
    qq_x, qq_y = 100, 50
    qq_w, qq_h = 200, 200
    draw.rectangle([qq_x, qq_y, qq_x + qq_w, qq_y + qq_h], 
                   fill=(240, 255, 240), outline=green, width=3)
    draw.text((qq_x + 20, qq_y + 10), "QUIZ QUESTION", fill=dark_gray, font=font_large)
    draw.text((qq_x + 20, qq_y + 40), "id (PK)", fill=dark_gray, font=font_small)
    draw.text((qq_x + 20, qq_y + 60), "question_text", fill=dark_gray, font=font_small)
    draw.text((qq_x + 20, qq_y + 80), "option_a, b, c, d", fill=dark_gray, font=font_small)
    draw.text((qq_x + 20, qq_y + 100), "correct_answer", fill=dark_gray, font=font_small)
    draw.text((qq_x + 20, qq_y + 120), "category, quiz_type", fill=dark_gray, font=font_small)
    draw.text((qq_x + 20, qq_y + 140), "explanation", fill=dark_gray, font=font_small)
    draw.text((qq_x + 20, qq_y + 160), "difficulty, time_limit", fill=dark_gray, font=font_small)
    draw.text((qq_x + 20, qq_y + 180), "created_at", fill=dark_gray, font=font_small)
    
    # LearningResource (Top Right)
    lr_x, lr_y = 900, 50
    lr_w, lr_h = 200, 150
    draw.rectangle([lr_x, lr_y, lr_x + lr_w, lr_y + lr_h], 
                   fill=(255, 255, 240), outline=(255, 212, 77), width=3)
    draw.text((lr_x + 20, lr_y + 10), "LEARNING RESOURCE", fill=dark_gray, font=font_large)
    draw.text((lr_x + 20, lr_y + 40), "id (PK)", fill=dark_gray, font=font_small)
    draw.text((lr_x + 20, lr_y + 60), "title, description", fill=dark_gray, font=font_small)
    draw.text((lr_x + 20, lr_y + 80), "url", fill=dark_gray, font=font_small)
    draw.text((lr_x + 20, lr_y + 100), "category, resource_type", fill=dark_gray, font=font_small)
    draw.text((lr_x + 20, lr_y + 120), "created_at", fill=dark_gray, font=font_small)
    
    # QuizType (Bottom Center)
    qt_x, qt_y = 500, 550
    qt_w, qt_h = 200, 120
    draw.rectangle([qt_x, qt_y, qt_x + qt_w, qt_y + qt_h], 
                   fill=(245, 240, 255), outline=(147, 51, 234), width=3)
    draw.text((qt_x + 20, qt_y + 10), "QUIZ TYPE", fill=dark_gray, font=font_large)
    draw.text((qt_x + 20, qt_y + 40), "id (PK)", fill=dark_gray, font=font_small)
    draw.text((qt_x + 20, qt_y + 60), "name, description", fill=dark_gray, font=font_small)
    draw.text((qt_x + 20, qt_y + 80), "time_limit, difficulty", fill=dark_gray, font=font_small)
    draw.text((qt_x + 20, qt_y + 100), "question_count", fill=dark_gray, font=font_small)
    
    # Relationships
    # User -> QuizAttempt (1 to Many)
    draw.line([user_x + 100, user_y + user_h, qa_x + 100, qa_y], fill=blue, width=2)
    draw.text((user_x + 50, user_y + user_h + 20), "1", fill=blue, font=font_medium)
    draw.text((qa_x + 50, qa_y - 20), "Many", fill=blue, font=font_medium)
    
    # User -> UserActivity (1 to Many)
    draw.line([user_x + user_w, user_y + 100, ua_x, ua_y + 50], fill=blue, width=2)
    draw.text((user_x + user_w + 10, user_y + 80), "1", fill=blue, font=font_medium)
    draw.text((ua_x - 50, ua_y + 30), "Many", fill=blue, font=font_medium)
    
    # QuizQuestion -> QuizAttempt (Many to Many through answers JSON)
    draw.line([qq_x + qq_w, qq_y + 100, qa_x, qa_y + 50], fill=green, width=2)
    draw.text((qq_x + qq_w + 10, qq_y + 80), "Many", fill=green, font=font_medium)
    draw.text((qa_x - 50, qa_y + 30), "Many", fill=green, font=font_medium)
    
    img_bytes = io.BytesIO()
    img.save(img_bytes, format='PNG')
    img_bytes.seek(0)
    return img_bytes

def clean_md_text(text):
    text = text.replace('**', '').replace('*', '')
    text = text.replace('`', '')
    text = text.replace('<br>', '\n')
    return text.strip()

def parse_md_and_add_to_doc(doc, md_file_path):
    with open(md_file_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    lines = content.split('\n')
    
    # Skip cover page, acknowledgement, declaration sections
    start_processing = False
    skip_sections = ['ACKNOWLEDGEMENT', 'DECLARATION', 'TABLE OF CONTENTS', 'LIST OF TABLES', 'LIST OF FIGURES']
    
    diagram_added = False
    er_added = False
    ml_pipeline_added = False
    
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        
        # Skip until we reach ABSTRACT or CHAPTER 1
        if not start_processing:
            if line.upper().startswith('# ABSTRACT') or '## ABSTRACT' in line.upper() or 'CHAPTER 1' in line.upper():
                start_processing = True
            else:
                i += 1
                continue
        
        # Skip certain sections
        if any(section in line.upper() for section in skip_sections):
            i += 1
            continue
        
        # Handle headers
        if line.startswith('# '):
            text = clean_md_text(line[2:])
            if text.upper() not in ["DIGITAL AWARENESS PLATFORM", "REPORT FILE OF PROJECT"]:
                doc.add_page_break()
                h = doc.add_heading(text, level=1)
                for r in h.runs:
                    set_font(r, font_size=16, bold=True)
        
        elif line.startswith('## '):
            text = clean_md_text(line[3:])
            if text.upper() not in skip_sections:
                h = doc.add_heading(text, level=2)
                for r in h.runs:
                    set_font(r, font_size=14, bold=True)
        
        elif line.startswith('### '):
            text = clean_md_text(line[4:])
            h = doc.add_heading(text, level=3)
            for r in h.runs:
                set_font(r, font_size=13, bold=True)
        
        elif line.startswith('#### '):
            text = clean_md_text(line[5:])
            h = doc.add_heading(text, level=3)
            for r in h.runs:
                set_font(r, font_size=12, bold=True)
        
        # Add diagrams at appropriate sections
        elif 'Figure 3.1' in line or ('System Architecture' in line and '3.1.2' in content[max(0, i-5):i+5]):
            if not diagram_added:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                img_bytes = create_system_architecture_diagram()
                run.add_picture(img_bytes, width=Inches(6))
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run("Figure 3.1: System Architecture of Digital Awareness Platform")
                set_font(run, bold=True)
                diagram_added = True
        
        elif 'Figure 3.3' in line or ('Machine Learning Pipeline' in line and '3.1.4' in content[max(0, i-5):i+5]):
            if not ml_pipeline_added:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                img_bytes = create_ml_pipeline_diagram()
                run.add_picture(img_bytes, width=Inches(6))
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run("Figure 3.3: Machine Learning Model Pipeline")
                set_font(run, bold=True)
                ml_pipeline_added = True
        
        elif 'Figure 4.1' in line or ('User Flow' in line and '4.5' in content[max(0, i-5):i+5]):
            # User flow is text-based, skip diagram
            pass
        
        # Handle code blocks
        elif line.startswith('```'):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            if code_lines:
                p = doc.add_paragraph()
                run = p.add_run('\n'.join(code_lines))
                run.font.name = 'Courier New'
                run.font.size = Pt(10)
        
        # Handle lists
        elif line.startswith('- ') or line.startswith('* '):
            text = clean_md_text(line[2:])
            p = doc.add_paragraph(text, style='List Bullet')
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            for r in p.runs:
                set_font(r)
        
        elif re.match(r'^\d+\.\s', line):
            text = clean_md_text(re.sub(r'^\d+\.\s', '', line))
            p = doc.add_paragraph(text, style='List Number')
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            for r in p.runs:
                set_font(r)
        
        # Handle tables (simple markdown tables)
        elif line.startswith('|') and '---' not in line:
            # Skip table rows for now (could parse if needed)
            pass
        
        # Handle regular paragraphs
        else:
            if line and not line.startswith('---'):
                text = clean_md_text(line)
                if text and len(text) > 3:  # Skip very short lines
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    run = p.add_run(text)
                    set_font(run)
        
        i += 1

def main():
    doc = Document()
    set_style(doc)
    
    # Add cover page
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("A PROJECT REPORT\nON\n")
    set_font(run, font_size=16, bold=True)
    run = p.add_run("\nDIGITAL AWARENESS PLATFORM\n")
    set_font(run, font_size=20, bold=True)
    run = p.add_run("\nSUBMITTED IN PARTIAL FULFILLMENT OF THE REQUIREMENTS\nFOR THE AWARD OF THE DEGREE OF\n\n")
    set_font(run, font_size=14)
    run = p.add_run("BACHELOR OF TECHNOLOGY\nIN\nCOMPUTER SCIENCE AND ENGINEERING\n\n")
    set_font(run, font_size=16, bold=True)
    run = p.add_run("By\n\n[Your Name]\n(Roll No: XXXXXXX)\n\n")
    set_font(run, font_size=14, bold=True)
    run = p.add_run("Under the Guidance of\n\n[Guide Name]\n[Designation]\n\n")
    set_font(run, font_size=14, bold=True)
    run = p.add_run("DEPARTMENT OF COMPUTER SCIENCE & ENGINEERING\n[INSTITUTE NAME]\n[CITY, STATE - PIN CODE]\n2024-2025")
    set_font(run, font_size=16, bold=True)
    
    doc.add_page_break()
    
    # Parse markdown and add content
    parse_md_and_add_to_doc(doc, 'Digital_Awareness_Platform_Project_Report_Final.md')
    
    # Save document
    output_path = 'Digital_Awareness_Platform_Project_Report_Final.docx'
    try:
        doc.save(output_path)
        print(f"Report generated successfully: {output_path}")
    except PermissionError:
        output_path = 'Digital_Awareness_Platform_Project_Report_Final_v3.docx'
        doc.save(output_path)
        print(f"Report generated successfully: {output_path}")

if __name__ == "__main__":
    main()

